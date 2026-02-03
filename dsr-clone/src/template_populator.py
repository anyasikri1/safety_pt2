"""Template populator: assemble a filled markdown document and .docx from parsed template sections.

Walks parsed template sections, resolves source references using ib_resolver,
and produces a single filled markdown document plus a .docx conversion.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from src.ib_resolver import resolve_sources
from src.models import TemplateSection
from src.utils import ensure_dir, logger


def _heading_level(section_id: str) -> int:
    """Determine the markdown heading level from a section_id's depth.

    - "1" -> ## (level 2)
    - "2.1" -> ### (level 3)
    - "2.1.1" -> #### (level 4)
    - Non-numeric ids (e.g. "Executive Summary") -> ## (level 2)
    - Cap at 6
    """
    parts = section_id.strip().split(".")
    try:
        # Verify all parts are numeric
        for p in parts:
            int(p)
        depth = len(parts)
        level = depth + 1  # depth 1 -> level 2, depth 2 -> level 3, etc.
        return min(level, 6)
    except ValueError:
        return 2


def assemble_markdown(
    template_sections: list[TemplateSection],
    ib_index: dict[str, str],
) -> str:
    """Build a single markdown document from template sections and resolved IB content.

    Starts with a top-level heading, then for each section:
    - Adds a heading at the appropriate level
    - If no required_sources: keeps the template body as-is
    - If 1 source: inserts ``*Source: {ref}*`` followed by content
    - If multiple sources: inserts each under a subheading ``### From {ref}``
    """
    lines: list[str] = ["# Filled Signal Assessment Report\n"]

    for section in template_sections:
        level = _heading_level(section.section_id)
        hashes = "#" * level
        lines.append(f"{hashes} {section.section_id} {section.title}\n")

        if not section.required_sources:
            # No sources -- keep template body as-is
            if section.body:
                lines.append(f"{section.body}\n")
        else:
            resolved = resolve_sources(section.required_sources, ib_index)
            if len(resolved) == 1:
                rs = resolved[0]
                lines.append(f"*Source: {rs.original_ref}*\n")
                lines.append(f"{rs.content}\n")
            else:
                for rs in resolved:
                    lines.append(f"### From {rs.original_ref}\n")
                    lines.append(f"{rs.content}\n")

    return "\n".join(lines)


def _markdown_to_docx(md_content: str, output_path: Path) -> None:
    """Convert markdown text to a .docx file using python-docx.

    Conversion rules:
    - Lines starting with ``#`` -> headings (level = number of hashes)
    - Lines wrapped in single ``*`` -> italic paragraph (source labels)
    - Lines starting with ``[MANUAL INPUT REQUIRED:`` or ``[CONTENT NOT FOUND:`` -> bold paragraph
    - Everything else -> normal paragraph
    - Skip empty lines
    """
    doc = Document()

    for line in md_content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Heading lines
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            hashes = heading_match.group(1)
            heading_text = heading_match.group(2)
            level = len(hashes)
            # python-docx heading levels are 0-based for Heading 1, but
            # add_heading(level=N) maps to "Heading N" style
            doc.add_heading(heading_text, level=min(level, 9))
            continue

        # Italic lines (wrapped in single asterisks)
        italic_match = re.match(r"^\*([^*]+)\*$", stripped)
        if italic_match:
            p = doc.add_paragraph()
            run = p.add_run(italic_match.group(1))
            run.italic = True
            continue

        # Bold lines (placeholders)
        if stripped.startswith("[MANUAL INPUT REQUIRED:") or stripped.startswith("[CONTENT NOT FOUND:"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            continue

        # Normal paragraph
        doc.add_paragraph(stripped)

    doc.save(str(output_path))


def write_filled_template(
    template_sections: list[TemplateSection],
    ib_index: dict[str, str],
    output_dir: str | Path,
) -> dict[str, Path]:
    """Write filled_template.md and filled_template.docx, return their paths.

    Returns:
        ``{"md": Path(...), "docx": Path(...)}``
    """
    output_dir = Path(output_dir)
    ensure_dir(output_dir)

    md_content = assemble_markdown(template_sections, ib_index)

    md_path = output_dir / "filled_template.md"
    md_path.write_text(md_content, encoding="utf-8")
    logger.info("Wrote filled markdown template to %s", md_path)

    docx_path = output_dir / "filled_template.docx"
    _markdown_to_docx(md_content, docx_path)
    logger.info("Wrote filled DOCX template to %s", docx_path)

    return {"md": md_path, "docx": docx_path}
